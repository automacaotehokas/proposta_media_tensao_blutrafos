# services/document/word_tables.py
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, List, Any
from .word_formatter import (
    set_row_height, set_column_widths, apply_paragraph_formatting,
    set_cell_shading, add_double_borders, set_table_left_indent
)

def determinar_eficiencia(perdas: str) -> str:
    """Determina a eficiência com base nas perdas"""
    if perdas == '5356-D':
        return "D"
    elif perdas == '5356-A':
        return "A"
    elif perdas == '1,2 %':
        return "1,2%"
    elif perdas == '1,0 %':
        return "1%"
    else:
        return "N/A"

def create_custom_table(doc: Document, itens_configurados: List[Dict], observacao: str) -> object:
    """Cria a tabela de preços customizada"""
    num_linhas = len(itens_configurados) + 2  # Uma linha para cada item + cabeçalho + total
    table = doc.add_table(rows=num_linhas, cols=10)  # Adicionando a coluna de IPI

    # Ajustar o alinhamento da tabela para a esquerda
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.left_indent = Cm(0)

    # Definir larguras fixas para as colunas
    col_widths = [Cm(1.1), Cm(1.25), Cm(2.2), Cm(1.0), Cm(2.7), Cm(1.0), 
                  Cm(1.75), Cm(2.63), Cm(2.63), Cm(1.15)]
    
    # Desabilitar o ajuste automático
    table.autofit = False
    set_column_widths(table, col_widths)

    # Cabeçalho
    header_row = table.rows[0]
    header_data = ["Item", "Qtde", "Potência", "K", "Tensões", "IP", "Perda", 
                  "Preço Uni. R$", "Preço Total R$", "IPI"]
    
    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (Títulos)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center')
        set_cell_shading(cell, '00543C')
        add_double_borders(cell)

    set_row_height(header_row, 1)

    # Preenchendo a tabela com os itens
    for idx, item in enumerate(itens_configurados, start=1):
        row = table.rows[idx]
        
        # Formatação da potência
        potencia = item["Potência"]
        potencia_texto = f"{potencia:,.1f}".replace('.', ',') + " kVA" if potencia % 1 != 0 else f"{int(potencia)} kVA"

        # Preenchimento das células
        cells_data = [
            str(idx),
            str(item["Quantidade"]),
            potencia_texto,
            str(item["Fator K"]),
            f"{item['Tensão Primária']}kV /{item['Tensão Secundária']} V",
            str(item["IP"]),
            str(item["Perdas"]),
            f"{item['Preço Unitário']:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            f"{item['Preço Total']:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            f"{item.get('IPI', '0')}%"
        ]

        for cell_idx, cell_text in enumerate(cells_data):
            cell = row.cells[cell_idx]
            cell.text = cell_text
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'Calibri Light (Títulos)'
            run.font.size = Pt(11)
            apply_paragraph_formatting(paragraph, alignment='center')
            add_double_borders(cell)

        set_row_height(row, 1.0)

    # Última linha - Valor Total
    total_row = table.rows[-1]
    total_row.cells[0].merge(total_row.cells[6])
    total_row.cells[7].merge(total_row.cells[9])
    
    total = sum(item['Preço Total'] for item in itens_configurados)
    total_row.cells[0].text = "Valor Total do Fornecimento:"
    total_row.cells[7].text = f"R$ {total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # Formatação da linha de total
    for idx in [0, 7]:
        cell = total_row.cells[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (Títulos)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center', space_before=Pt(0))
        set_cell_shading(cell, '00543C')
        add_double_borders(cell)

    set_row_height(total_row, 0.6)

    # Adicionar observação
    obs_row = table.add_row()
    obs_row.cells[0].merge(obs_row.cells[9])
    obs_cell = obs_row.cells[0]
    obs_cell.text = f"Obs.: {observacao}"
    obs_paragraph = obs_cell.paragraphs[0]
    obs_run = obs_paragraph.runs[0]
    obs_run.font.name = 'Calibri Light (Títulos)'
    obs_run.font.size = Pt(11)
    apply_paragraph_formatting(obs_paragraph, alignment='left', space_before=Pt(0))
    add_double_borders(obs_cell)

    return table

def create_custom_table_escopo(doc: Document, itens_configurados: List[Dict]) -> object:
    """Cria a tabela de escopo"""
    table = doc.add_table(rows=len(itens_configurados) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_left_indent(table, 0)
    table.left_indent = Cm(0)
    table.autofit = False

    # Definir larguras das colunas
    col_widths = [Cm(1.5), Cm(15.0)]
    set_column_widths(table, col_widths)

    # Cabeçalho
    header_row = table.rows[0]
    header_data = ["Item", "Escopo do Fornecimento:"]

    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (T)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center')
        set_row_height(header_row, 1)
        set_cell_shading(cell, '00543C')
        add_double_borders(cell)

    # Preenchendo itens
    for idx, item in enumerate(itens_configurados, start=1):
        row = table.rows[idx]
        
        # Coluna "Item"
        row.cells[0].text = str(idx)
        apply_paragraph_formatting(row.cells[0].paragraphs[0], alignment='center')
        add_double_borders(row.cells[0])

        # Preparar dados para o escopo
        classe_tensao = item.get('classe_tensao', '').replace('kV', '').strip()
        tensao_secundaria = item.get('Tensão Secundária', 'N/A').replace('kV', '').strip()
        eficiencia = determinar_eficiencia(item['Perdas'])
        NBI = item.get('NBI', 'N/A')
        
        # Formatar potência
        potencia = item.get('Potência', 'N/A')
        if isinstance(potencia, (int, float)):
            if float(potencia).is_integer():  # Verifica se é um número inteiro
                potencia_formatada = f"{int(potencia)} kVA"
            else:
                potencia_formatada = f"{potencia:.1f}".replace('.', ',') + " kVA"
        else:
            potencia_formatada = potencia

        # Calcular tensão secundária
        tensao_secundaria_str = item.get('Tensão Secundária', '0')
        try:
            tensao_secundaria_float = float(tensao_secundaria_str)
            tensao_calculada = tensao_secundaria_float / 1.73
            tensao_calculada_arredondada = round(tensao_calculada)
            tensao_secundaria_arredondada = round(tensao_secundaria_float)
            tensao_secundaria_texto = f"{tensao_secundaria_arredondada}/{tensao_calculada_arredondada}V"
        except ValueError:
            tensao_secundaria_texto = f"{tensao_secundaria_str}V"

        escopo_text = (
            f"Transformador Trifásico **isolado a seco**, Classe de tensão **{classe_tensao}/1,1kV**, "
            f"Marca e Fabricação Blutrafos, Potência: **{potencia_formatada}**, Fator: **K={item.get('Fator K', 'N/A')}**, "
            f"Tensão **Primária**: **{item.get('Tensão Primária', 'N/A')}kV**, Derivações: **{item.get('Derivações', 'N/A')}**, "
            f"Tensão **Secundária**: **{tensao_secundaria_texto}**, Grupo de Ligação: **Dyn-1**, "
            f"Frequência: **60Hz**, NBI: **{NBI}**, Classe de Temperatura: F (155ºC), "
            f"Elevação Temperatura média dos enrolamentos: **100ºC**, Materiais dos enrolamentos: **Alumínio**, "
            f"Altitude de Instalação: **≤1000m**, Temperatura ambiente máxima: 40°C, "
            f"Alta tensão Encapsulado em Resina Epóxi à Vácuo, Regime de Serviço: Contínuo, "
            f"Tipo de Refrigeração: **AN** e Grau de Proteção: **IP-{item.get('IP', 'N/A')}**, "
            f"Demais características cfe. Norma ABNT-NBR 5356/11 - Eficiência **“{eficiencia}”** e acessórios abaixo."
        )

        # Aplicar o texto com formatação
        escopo_paragraph = row.cells[1].paragraphs[0]
        escopo_paragraph.text = ""
        escopo_parts = escopo_text.split("**")
        
        for i, part in enumerate(escopo_parts):
            run = escopo_paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True
            run.font.name = 'Calibri Light (Título)'
            run.font.size = Pt(10)

        escopo_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        escopo_paragraph.paragraph_format.space_after = Pt(2)
        add_double_borders(row.cells[1])

    return table

def inserir_tabelas_word(doc: Document, itens_configurados: List[Dict], 
                        observacao: str, replacements: Dict) -> Document:
    """Insere as tabelas no documento e realiza substituições de texto"""
    # Substituir texto no documento
    substituir_texto_documento(doc, replacements)

    # Inserir tabela de preços
    for i, paragraph in enumerate(doc.paragraphs):
        if "Quadro de Preços" in paragraph.text:
            table = create_custom_table(doc, itens_configurados, observacao)
            doc.paragraphs[i+1]._element.addnext(table._element)
            break

    # Inserir tabela de escopo
    for i, paragraph in enumerate(doc.paragraphs):
        if "Escopo de Fornecimento" in paragraph.text:
            table_escopo = create_custom_table_escopo(doc, itens_configurados)
            doc.paragraphs[i+1]._element.addnext(table_escopo._element)
            break

    return doc

def substituir_texto_documento(doc: Document, replacements: Dict[str, str]):
    """Substitui textos no documento inteiro"""
    def remove_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    # Substituir em parágrafos
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                if old_text == "{{IP}}" and not new_text.strip():
                    remove_paragraph(paragraph)
                    break
                else:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

    # Substituir em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            if old_text == "{{IP}}" and not new_text.strip():
                                remove_paragraph(paragraph)
                                break
                            else:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                         # Substituir no cabeçalho
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    if old_text == "{{IP}}" and not new_text.strip():
                        remove_paragraph(paragraph)
                        break
                    else:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)

def get_table_width(doc: Document) -> float:
    """Retorna a largura disponível para tabelas no documento"""
    section = doc.sections[0]
    page_width = section.page_width
    margin_left = section.left_margin
    margin_right = section.right_margin
    return page_width - margin_left - margin_right

def format_currency(value: float) -> str:
    """Formata valor monetário no padrão brasileiro"""
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def format_potencia(potencia: Any) -> str:
    """Formata o valor da potência"""
    if isinstance(potencia, (int, float)):
        if potencia % 1 == 0:
            return f"{int(potencia)} kVA"
        return f"{potencia:.1f}".replace('.', ',') + " kVA"
    return str(potencia)

def calcular_tensao_secundaria(tensao_str: str) -> str:
    """Calcula e formata a tensão secundária"""
    try:
        tensao = float(tensao_str)
        tensao_fase = tensao / 1.73
        return f"{round(tensao)}/{round(tensao_fase)}V"
    except (ValueError, TypeError):
        return f"{tensao_str}V"
    

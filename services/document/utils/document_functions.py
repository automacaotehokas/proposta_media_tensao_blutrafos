from typing import Dict
import logging
import docx2pdf
import tempfile
from docx.shared import Pt,RGBColor,Inches
from docx import Document
import os
import dotenv
from shareplum import Site, Office365
from shareplum.site import Version
from dotenv import load_dotenv
import streamlit as st
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Importações dos serviços MT
from services.document.mt import (
    pdf_service_mt,
    word_service_mt,
    word_formatter_mt,
    word_tables_mt
)

from services.document.bt import (
    pdf_service_bt,
    word_service_bt,
    word_formatter_bt,
    word_tables_bt
)

from pages.pagamento_entrega.components import ComponentsPagamentoEntrega
import tempfile

from services.sharepoint.sharepoint_service import SharePoint
from pages.pagamento_entrega.components import ComponentsPagamentoEntrega
import logging

# Configuração do logger
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)


logger = logging.getLogger(__name__)



def inserir_desvios(doc):
    """
    Insere desvios no documento Word usando o estilo Bullet nativo do Word.
    
    Args:
        doc: Documento Word aberto com python-docx
    """
    # Encontra o parágrafo que contém o texto de referência
    texto_referencia = "repassadas de forma a objetivar o equilíbrio do contrato."
    paragrafo_inicial = None
    
    for paragraph in doc.paragraphs:
        if texto_referencia in paragraph.text:
            paragrafo_inicial = paragraph
            break
    
    if not paragrafo_inicial:
        raise ValueError("Parágrafo de referência não encontrado")

    ultimo_paragrafo = paragrafo_inicial
    
    if 'desvios' in st.session_state and st.session_state['desvios']:
        for desvio in st.session_state['desvios']:
            # Usa diretamente o estilo Bullet
            p_desvio = doc.add_paragraph(style='Bullet')
            
            # Adiciona o checkmark e o texto
            run_desvio = p_desvio.add_run(desvio['texto'])
            run_desvio.font.size = Pt(11)
            
            # Ajusta a indentação se necessário
            p_desvio.paragraph_format.left_indent = Inches(0.5)
            
            # Posiciona o novo parágrafo
            ultimo_paragrafo._element.addnext(p_desvio._element)
            ultimo_paragrafo = p_desvio


import re
from docx.shared import RGBColor, Pt

import re
from docx.shared import RGBColor, Pt

def highlight_work_days(paragraph, input_text):
    """
    Highlights numbers followed by variations of 'dias uteis' in red color
    and 'Prazo de fabricação' in bold.
    
    Args:
        paragraph: A python-docx paragraph object where the text will be inserted
        input_text: The input text to process
    """
    # Clear existing content from paragraph
    paragraph.clear()
    
    # Patterns
    dias_pattern = r'(\d+)(\s*(?:dias?\s*[uú]te[ie]s?))'
    prazo_pattern = r'(prazo\s+de\s+fabrica[çc][aã]o)'
    
    # Function to convert text to lowercase for comparison
    def normalize_text(text):
        return text.lower().replace('ú', 'u').replace('é', 'e').replace('í', 'i').replace('ç', 'c').replace('ã', 'a')
    
    # Combine both patterns to process text in order
    combined_matches = []
    
    # Find all matches for dias uteis
    for match in re.finditer(dias_pattern, input_text, re.IGNORECASE):
        combined_matches.append(('dias', match))
    
    # Find all matches for prazo de fabricação
    for match in re.finditer(prazo_pattern, input_text, re.IGNORECASE):
        combined_matches.append(('prazo', match))
    
    # Sort matches by their position in text
    combined_matches.sort(key=lambda x: x[1].start())
    
    # Process each match in order
    last_end = 0
    for match_type, match in combined_matches:
        # Add text before the match
        if match.start() > last_end:
            run = paragraph.add_run(input_text[last_end:match.start()])
            run.font.name = 'Calibri Light'
            run.font.size = Pt(11)
        
        if match_type == 'dias':
            # Add the number in red
            number = match.group(1)
            run = paragraph.add_run(number)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            run.font.name = 'Calibri Light'
            run.font.size = Pt(11)
            
            # Add "dias uteis" in red
            dias_uteis = match.group(2)
            run = paragraph.add_run(dias_uteis)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            run.font.name = 'Calibri Light'
            run.font.size = Pt(11)
        
        elif match_type == 'prazo':
            # Add "Prazo de fabricação" in bold, preserving original case
            prazo_text = match.group(0)
            run = paragraph.add_run(prazo_text)
            run.font.name = 'Calibri Light'
            run.font.size = Pt(11)
            run.bold = True
        
        last_end = match.end()
    
    # Add any remaining text after the last match
    if last_end < len(input_text):
        run = paragraph.add_run(input_text[last_end:])
        run.font.name = 'Calibri Light'
        run.font.size = Pt(11)

# Example usage:
# doc = Document()
# paragraph = doc.add_paragraph()
# highlight_work_days(paragraph, "Prazo de Fabricação: 60 dias úteis após PRAZO DE FABRICAÇÃO.")
# Example usage:
# doc = Document()
# paragraph = doc.add_paragraph()
# highlight_work_days(paragraph, "Prazo de entrega: 60 dias úteis após aprovação.")

def inserir_prazo_entrega(doc):
    """
    Insere os prazos de entrega no documento Word, processando marcações de negrito
    indicadas por asteriscos duplos (**).
    """
    # Função auxiliar para processar texto com marcações de negrito
    def processar_texto_com_negrito(paragraph, texto, valor_colorido=None, incluir_dias_uteis=False):
        """
        Processa texto com marcações de negrito e aplica cor vermelha ao valor e 'dias úteis'.
        
        Args:
            paragraph: O parágrafo onde o texto será inserido
            texto: O texto com marcações de negrito
            valor_colorido: O valor que deve aparecer em vermelho
            incluir_dias_uteis: Se True, também colore 'dias úteis' em vermelho
        """
        partes = texto.split("**")
        
        for i, parte in enumerate(partes):
            if valor_colorido is not None and str(valor_colorido) in parte:
                # Primeiro, dividimos pelo valor numérico
                antes, valor, resto = parte.partition(str(valor_colorido))
                
                # Adicionamos o texto antes do valor
                if antes:
                    run = paragraph.add_run(antes)
                    run.font.name = 'Calibri Light'
                    run.font.size = Pt(11)
                    if i % 2 == 1:
                        run.bold = True
                
                # Adicionamos o valor em vermelho
                if valor:
                    run = paragraph.add_run(valor)
                    run.font.name = 'Calibri Light'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Vermelho
                    if i % 2 == 1:
                        run.bold = True
                
                # Se precisamos colorir "dias úteis", vamos procurar por isso no resto do texto
                if incluir_dias_uteis and "dias úteis" in resto:
                    # Dividimos o resto do texto em: antes dos "dias úteis", "dias úteis", e depois
                    antes_dias, dias_uteis, depois = resto.partition("dias úteis")
                    
                    # Adiciona o texto antes de "dias úteis"
                    if antes_dias:
                        run = paragraph.add_run(antes_dias)
                        run.font.name = 'Calibri Light'
                        run.font.size = Pt(11)
                        if i % 2 == 1:
                            run.bold = True
                    
                    # Adiciona "dias úteis" em vermelho
                    run = paragraph.add_run("dias úteis")
                    run.font.name = 'Calibri Light'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Vermelho
                    if i % 2 == 1:
                        run.bold = True
                    
                    # Adiciona o texto restante
                    if depois:
                        run = paragraph.add_run(depois)
                        run.font.name = 'Calibri Light'
                        run.font.size = Pt(11)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    # Se não precisamos colorir "dias úteis", adiciona o resto normalmente
                    if resto:
                        run = paragraph.add_run(resto)
                        run.font.name = 'Calibri Light'
                        run.font.size = Pt(11)
                        if i % 2 == 1:
                            run.bold = True
            else:
                # Parte sem valor para colorir
                run = paragraph.add_run(parte)
                run.font.name = 'Calibri Light'
                run.font.size = Pt(11)
                if i % 2 == 1:
                    run.bold = True

    # Agora, ao usar a função para o prazo de desenho:


    # Na parte do código onde processamos o prazo de desenho:

    # Localiza o parágrafo inicial
    paragrafo_inicial = None
    for paragraph in doc.paragraphs:
        if "A partir destes eventos, consideramos os seguintes prazos:" in paragraph.text:
            paragrafo_inicial = paragraph
            break
    
    if not paragrafo_inicial:
        raise ValueError("Texto inicial dos prazos não encontrado no documento")

    ultimo_paragrafo = paragrafo_inicial

    # Processar prazo de desenho
    if 'prazo_desenho' in st.session_state['prazo_entrega']:
        valor_desenho = st.session_state['prazo_entrega']['prazo_desenho']
        p_desenho = doc.add_paragraph(style='Bullet')
        texto_desenho = f"**Desenhos para aprovação:** Até {valor_desenho} dias úteis contados a partir da data de efetivação das etapas **a** e **b**, listadas acima."
        
        # Note o novo parâmetro incluir_dias_uteis=True
        processar_texto_com_negrito(p_desenho, texto_desenho, valor_desenho, incluir_dias_uteis=True)
        ultimo_paragrafo._element.addnext(p_desenho._element)
        ultimo_paragrafo = p_desenho

    # Processar prazo de cliente
    if 'prazo_cliente' in st.session_state['prazo_entrega']:
        valor_cliente = st.session_state['prazo_entrega']['prazo_cliente']
        p_cliente = doc.add_paragraph(style='Bullet')
        texto_cliente = f"**Prazo para aprovação dos desenhos pelo cliente:** Até {valor_cliente} dias úteis contados a partir da data de envio dos desenhos para aprovação. Se o tempo de aprovação for maior que informado, o prazo de entrega será obrigatoriamente renegociado"
        
        # Usa a função auxiliar para processar o texto
        processar_texto_com_negrito(p_cliente, texto_cliente,valor_cliente, incluir_dias_uteis=True)
        ultimo_paragrafo._element.addnext(p_cliente._element)
        ultimo_paragrafo = p_cliente

    # Processar prazo de fabricação
    if 'prazo_entrega_global' in st.session_state and 'prazo_fabricacao' in st.session_state['prazo_entrega_global']:
        try:
            valor = st.session_state['prazo_entrega_global']['prazo_fabricacao']
            evento = "Aprovação dos Desenhos"
            
            if valor:
                p_produto = doc.add_paragraph(style='Bullet')
                texto_produto = f"{valor}"
                
                # Usa a função auxiliar para processar o texto
                highlight_work_days(p_produto,texto_produto)
                ultimo_paragrafo._element.addnext(p_produto._element)
                ultimo_paragrafo = p_produto
                logger.debug(f"Prazo de fabricação inserido: {valor} dias, evento: {evento}")
        except Exception as e:
            logger.error(f"Erro ao inserir prazo de fabricação: {str(e)}")



def inserir_eventos_pagamento(doc, eventos_pagamento, produtos_configurados):
    """
    Insere eventos de pagamento no documento Word.
    
    Args:
        doc: Documento Word
        eventos_pagamento: Lista de eventos do session_state
        produtos_configurados: Dicionário indicando quais produtos estão configurados
    """
    logger.info("Iniciando inserção de texto de pagamento")
    
    def extrair_eventos_do_session_state():
        """
        Extrai e valida os eventos de pagamento do session_state.
        Retorna uma lista de eventos validados no formato correto.
        """
        try:
            # Se eventos_pagamento é uma lista direta do session_state
            if isinstance(eventos_pagamento, list):
                return [
                    {
                        'percentual': float(evento['percentual']),
                        'dias': evento['dias'],
                        'evento': str(evento['evento'])
                    }
                    for evento in eventos_pagamento
                    if isinstance(evento, dict) and 
                    all(key in evento for key in ['percentual', 'dias', 'evento'])
                ]
            
            # Se recebemos o session_state completo
            if hasattr(eventos_pagamento, 'eventos_pagamento'):
                eventos_temp = eventos_pagamento.eventos_pagamento
                if isinstance(eventos_temp, list):
                    return [
                        {
                            'percentual': float(evento['percentual']),
                            'dias': evento['dias'],
                            'evento': str(evento['evento'])
                        }
                        for evento in eventos_temp
                        if isinstance(evento, dict) and 
                        all(key in evento for key in ['percentual', 'dias', 'evento'])
                    ]
            
            logger.warning("Formato de eventos não reconhecido")
            return []
            
        except Exception as e:
            logger.error(f"Erro ao extrair eventos: {str(e)}")
            return []

    try:
        # Extrair e validar os eventos
        eventos_validados = extrair_eventos_do_session_state()
        
        if not eventos_validados:
            logger.warning("Nenhum evento válido para processar")
            return
            
        # Encontrar o parágrafo "Condições de Pagamento"
        index = None
        for i, paragraph in enumerate(doc.paragraphs):
            if "Condições de Pagamento" in paragraph.text:
                index = i
                break
        
        if index is None:
            logger.error("Não foi encontrado o parágrafo 'Condições de Pagamento'")
            raise ValueError("Não foi encontrado o parágrafo 'Condições de Pagamento'")
        
        # Verifica se há produtos configurados
        tem_mt = produtos_configurados.get('mt', False)
        tem_bt = produtos_configurados.get('bt', False)
        
        if not (tem_mt or tem_bt):
            logger.warning("Nenhum produto MT ou BT configurado")
            return
            
        # Ordena os eventos por percentual (maior para menor)
        eventos_ordenados = sorted(eventos_validados, key=lambda x: float(x['percentual']), reverse=True)
        
        # Processa cada evento
        index_atual = index + 1
        for evento in eventos_ordenados:
            # Define o texto dos dias baseado no tipo de evento
            dias = evento['dias']
            evento_texto = evento['evento']
            
            # Mapeamento para formatar o texto dos dias
            dias_formato = {
                "Aprovação dos Desenhos": lambda d: 'com a' if d == "0" else f"{d} dias da",
                "Faturamento (Mediante aprovação financeira)": lambda d: 'após o' if d == "0" else f" a {d} dias do",
                "Contra aviso de pronto p/ embarque": lambda d: '' if d == "0"  else f"a {d} dias do",
                "TAF": lambda d: 'com o' if d == "0"  else f"{d} dias do",
                "Pedido": lambda d: 'com o' if d == "0" else f"{d} dias do",
                "Entrega do equipamento": lambda d: 'com o' if d == 0 else f"{d} dias do"
            }
            
            # Obtém o formato específico ou usa o padrão
            formato = dias_formato.get(evento_texto, lambda d: 'com o' if d == 0 else f"{d} dias do")
            dias_texto = formato(dias)
            
            # Cria e formata o parágrafo do evento
            p_evento = doc.add_paragraph(style='Bullet')
            p_evento.add_run(f"{int(evento['percentual'])}% - {dias_texto} {evento_texto};")
            p_evento.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Insere no documento
            doc.paragraphs[index_atual]._element.addnext(p_evento._element)
            index_atual += 1
            
        # Adiciona espaço após os eventos
        p_espaco = doc.add_paragraph()
        doc.paragraphs[index_atual]._element.addnext(p_espaco._element)
        
        logger.info("Inserção de eventos de pagamento concluída com sucesso")
        
    except Exception as e:
        logger.error(f"Erro durante a inserção dos eventos: {str(e)}")
        raise

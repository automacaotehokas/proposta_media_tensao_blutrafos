from docx import Document  # Importação específica
from typing import Dict, List, Optional
import logging
from io import BytesIO
# Importações de serviços
from services.document.mt import word_tables_mt
from services.document.bt import word_tables_bt
from services.document.mt.word_tables_mt import substituir_texto_documento
import streamlit as st
from ..utils.document_functions import inserir_desvios,inserir_eventos_pagamento,inserir_prazo_entrega
from pages.pagamento_entrega.components import ComponentsPagamentoEntrega
from ..mt.test import formatar_numero_inteiro_ou_decimal



logger = logging.getLogger(__name__)

def formatar_numero_brasileiro(valor):
    """
    Converte um número para o formato brasileiro de moeda.
    Exemplo: 10000000.50 -> '10.000.000,50'
    """
    try:
        # Converte para float para garantir formato correto
        valor = float(valor)
        # Usa format para separar milhares com ponto e usar vírgula para decimais
        return f"{valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except (ValueError, TypeError):
        return '0,00'

def inserir_tabelas_word(doc: Document, itens_configurados: List[Dict], 
                        observacao: str, replacements: Dict) -> Document:
    """Insere as tabelas no documento e realiza substituições de texto"""
    # Substituir texto no documento
    substituir_texto_documento(doc, replacements)

    # Inserir tabela de preços
    for i, paragraph in enumerate(doc.paragraphs):
        if "Quadro de Preços" in paragraph.text:
            table = word_tables_mt.create_custom_table(doc, itens_configurados, observacao)
            doc.paragraphs[i+1]._element.addnext(table._element)
            break

    # Inserir tabela de escopo
    # for i, paragraph in enumerate(doc.paragraphs):
    #     if "Escopo de Fornecimento" in paragraph.text:
    #         table_escopo = create_custom_table_escopo(doc, itens_configurados)
    #         doc.paragraphs[i+1]._element.addnext(table_escopo._element)
    #         break

    return doc

def inserir_tabelas_separadas(
    doc: Document, 
    itens_mt: Optional[List[Dict]] = None, 
    itens_bt: Optional[List[Dict]] = None, 
    observacao: str = "",
    replacements: Optional[Dict] = None
) -> Document:
    """
    Insere tabelas de MT e BT separadamente no documento Word
    """
    # Validar e preparar argumentos
    itens_mt = itens_mt or []
    itens_bt = itens_bt or []
    replacements = replacements or {}
    
    # Validar entrada
    if not itens_mt and not itens_bt:
        logger.warning("Nenhum item de MT ou BT fornecido. Nenhuma tabela será inserida.")
        return doc

    # Controle de inserção das tabelas de preços
    mt_price_inserted = False
    bt_price_inserted = False

    # Procurar marcadores e inserir tabelas de preços
    for i, paragraph in enumerate(doc.paragraphs):
        # Inserir tabela de preços MT
        if not mt_price_inserted and "{{ QUADRO_PRECOS_MT}}" in paragraph.text and itens_mt:
            logger.info(f"Processando itens MT: {len(itens_mt)} itens")
            try:
                # Criar tabela de MT
                table = word_tables_mt.create_custom_table(doc, itens_mt, observacao)
                doc.paragraphs[i + 1]._element.addnext(table._element)
                paragraph.text = paragraph.text.replace("{{ QUADRO_PRECOS_MT}}", "")
                mt_price_inserted = True
                logger.info("Tabela MT inserida com sucesso")
            except Exception as e:
                logger.error(f"Erro ao criar tabela MT: {str(e)}", exc_info=True)

        # Inserir tabela de preços BT
        if not bt_price_inserted and "{{ QUADRO_PRECOS_BT}}" in paragraph.text and itens_bt:
            logger.info(f"Processando itens BT: {len(itens_bt)} itens")
            try:
                # Criar tabela de BT
                table = word_tables_bt.create_custom_table(doc, itens_bt, observacao)
                doc.paragraphs[i + 1]._element.addnext(table._element)
                paragraph.text = paragraph.text.replace("{{ QUADRO_PRECOS_BT}}", "")
                bt_price_inserted = True
                logger.info("Tabela BT inserida com sucesso")
            except Exception as e:
                logger.error(f"Erro ao criar tabela BT: {str(e)}", exc_info=True)

        # Se ambas as tabelas necessárias foram inseridas, podemos parar
        if (mt_price_inserted or not itens_mt) and (bt_price_inserted or not itens_bt):
            break

    # Avisos para tabelas de preços não inseridas
    if not mt_price_inserted and itens_mt:
        logger.warning("Marcador de quadro de preços MT não encontrado no documento")
    if not bt_price_inserted and itens_bt:
        logger.warning("Marcador de quadro de preços BT não encontrado no documento")

    # Inserção das tabelas de escopo (mantendo a lógica existente)
    if itens_mt or itens_bt:
        mt_inserted = False
        bt_inserted = False
        
        for i, paragraph in enumerate(doc.paragraphs):
            # Handle MT escopo table
            if not mt_inserted and "{{ ESCOPO_MT}}" in paragraph.text and itens_mt:
                logger.info("Inserindo tabela de escopo MT")
                if hasattr(word_tables_mt, 'create_custom_table_escopo'):
                    tabela_escopo_mt = word_tables_mt.create_custom_table_escopo(doc, itens_mt)
                    doc.paragraphs[i + 1]._element.addnext(tabela_escopo_mt._element)
                paragraph.text = paragraph.text.replace("{{ ESCOPO_MT}}", "")
                mt_inserted = True

            # Handle BT escopo table
            if not bt_inserted and "{{ ESCOPO_BT}}" in paragraph.text and itens_bt:
                logger.info("Inserindo tabela de escopo BT")
                if hasattr(word_tables_bt, 'create_custom_table_escopo'):
                    tabela_escopo_bt = word_tables_bt.create_custom_table_escopo(doc, itens_bt)
                    doc.paragraphs[i + 1]._element.addnext(tabela_escopo_bt._element)
                paragraph.text = paragraph.text.replace("{{ ESCOPO_BT}}", "")
                bt_inserted = True

            if (mt_inserted or not itens_mt) and (bt_inserted or not itens_bt):
                break

        if not mt_inserted and itens_mt:
            logger.warning("Marcador de escopo MT não encontrado no documento")
        if not bt_inserted and itens_bt:
            logger.warning("Marcador de escopo BT não encontrado no documento")

    return doc
    
    


logger = logging.getLogger(__name__)


def substituir_texto_documento_1(doc, replacements):
    """
    Substitui texto no documento Word mantendo sua estrutura original.
    
    Args:
        doc: Documento Word (objeto Document)
        replacements: Dicionário com as substituições a serem feitas
        
    Returns:
        Document: Documento Word com as substituições realizadas
    """
    if not doc:
        raise ValueError("Documento não pode ser None")
        
    logger.info("Iniciando substituições com: %s", replacements)
    
    # Função auxiliar para substituir texto em um parágrafo
    def substituir_em_paragrafo(paragraph):
        # Obtém todo o texto do parágrafo
        texto_completo = ''.join(run.text for run in paragraph.runs)
        texto_modificado = texto_completo
        
        # Realiza todas as substituições necessárias
        for old_text, new_text in replacements.items():
            if old_text in texto_completo:
                texto_modificado = texto_modificado.replace(old_text, str(new_text))
                logger.info(f"Substituindo '{old_text}' por '{new_text}'")
                logger.info(f"Texto antes: {texto_completo}")
                logger.info(f"Texto depois: {texto_modificado}")
        
        # Se houve alteração, atualiza o texto no documento
        if texto_modificado != texto_completo:
            # Limpa todos os runs
            for run in paragraph.runs:
                run.text = ""
            # Coloca todo o texto modificado
            paragraph.add_run(texto_modificado)
    
    # Processa cada tipo de conteúdo do documento
    # 1. Corpo do documento
    logger.info("Processando corpo do documento")
    for paragraph in doc.paragraphs:
        substituir_em_paragrafo(paragraph)
    
    # 2. Cabeçalhos e rodapés
    logger.info("Processando cabeçalhos e rodapés")
    for section in doc.sections:
        # Cabeçalho
        header = section.header
        for paragraph in header.paragraphs:
            substituir_em_paragrafo(paragraph)
        
        # Rodapé
        footer = section.footer
        for paragraph in footer.paragraphs:
            substituir_em_paragrafo(paragraph)
    
    # 3. Tabelas
    logger.info("Processando tabelas")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_em_paragrafo(paragraph)
    
    # Verifica o resultado final
    logger.info("Verificando resultado final")
    for paragraph in doc.paragraphs:
        logger.info(f"Parágrafo após substituições: {paragraph.text}")
    
    # Verifica também os cabeçalhos, rodapés e tabelas
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            logger.info(f"Cabeçalho após substituições: {paragraph.text}")
        for paragraph in section.footer.paragraphs:
            logger.info(f"Rodapé após substituições: {paragraph.text}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    logger.info(f"Célula após substituições: {paragraph.text}")
    
    return doc
                            
def gerar_documento(template_path, dados_iniciais, impostos, itens_mt, itens_bt):
    try:
        logger.info("Iniciando geração de documento unificado")
        
        # Garantir que o template_path existe e está correto
        if not template_path:
            raise ValueError("Template path não pode ser vazio")

        # Carregar o documento
        try:
            doc = Document(template_path)
        except Exception as e:
            logger.error(f"Erro ao carregar template: {str(e)}")
            raise ValueError(f"Falha ao carregar template: {str(e)}")

        # Criar replacements
        replacements = {
            '{{CLIENTE}}': str(dados_iniciais.get('cliente', '')),
            '{{NOMECLIENTE}}': str(dados_iniciais.get('nomeCliente', '')),
            '{{FONE}}': str(dados_iniciais.get('fone', '')),
            '{{EMAIL}}': str(dados_iniciais.get('email', '')),
            '{{BT}}': str(dados_iniciais.get('bt', '')),
            '{{OBRA}}': str(dados_iniciais.get('obra', ' ')),
            '{{DIA}}': str(dados_iniciais.get('dia', '')),
            '{{MES}}': str(dados_iniciais.get('mes', '')),
            '{{ANO}}': str(dados_iniciais.get('ano', '')),
            '{{REV}}': str(dados_iniciais.get('rev', '')),
            '{{LOCAL}}': str(dados_iniciais.get('local_frete', '')),
            '{{LOCALFRETE}}': str(impostos.get('local_entrega', '')),
            '{{ICMS}}': f"{formatar_numero_inteiro_ou_decimal(impostos.get('icms', 0)):}",
            '{{DIFAL}}': f"{impostos.get('difal', 0):.1f}" if impostos.get('difal', 0) > 0 else '',
            '{{IP}}': '',
            '{obra}': '' if not dados_iniciais.get('obra', '').strip() else 'Obra',
            '{{VALOR_TOTAL}}': formatar_numero_brasileiro(st.session_state.get('valor_totalizado', 0)),
            '{{TRANSPORTE}}': f'CIP - {str(dados_iniciais.get('local_frete', ''))}, sobre o veículo transportador (descarga não inclusa)' if impostos.get('tipo_frete', '') == 'CIP' else f'FOB - {str(dados_iniciais.get('local_frete', ''))}, sobre o veículo transportador (descarga não inclusa)',
            '{{RESPONSAVEL}}': st.session_state.get('usuario', '')
        }

        # Carregar itens e produtos configurados
        itens = st.session_state.get('itens', {})
        produtos_configurados = ComponentsPagamentoEntrega.carregar_tipo_produto(itens)

        # Aplicar substituições
        logger.info("Aplicando substituições de texto")
        doc = substituir_texto_documento(doc, replacements)

        # Inserir tabelas
        logger.info("Inserindo tabelas")
        doc = inserir_tabelas_separadas(doc, itens_mt, itens_bt, "", replacements)

        # Inserir eventos e desvios
        inserir_eventos_pagamento(doc, st.session_state, produtos_configurados)
        inserir_desvios(doc)
        inserir_prazo_entrega(doc)

        # Salvar o documento em um buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Verificar o documento final
        logger.info("Verificando documento final")
        Document(BytesIO(buffer.getvalue()))  # Verificação final

        return buffer

    except Exception as e:
        logger.error(f"Erro na geração do documento: {str(e)}", exc_info=True)
        raise Exception(f"Erro ao gerar documentos: {str(e)}")
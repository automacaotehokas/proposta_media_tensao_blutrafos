from docx import Document
from typing import Dict, List
from io import BytesIO
import logging
from .word_tables_unificado import  inserir_tabelas_separadas

logger = logging.getLogger(__name__)

# def substituir_texto_documento(doc, replacements):
#     if not doc:
#         raise ValueError("Documento não pode ser None")

#     logger.info(f"Substituições a serem feitas: {replacements}")

#     def substituir_em_paragrafo(paragraph):
#         texto_completo = ''.join(run.text for run in paragraph.runs)
#         texto_modificado = texto_completo

#         for old_text, new_text in replacements.items():
#             if old_text in texto_completo:
#                 texto_modificado = texto_modificado.replace(old_text, str(new_text))

#         if texto_modificado != texto_completo:
#             # Limpa todos os runs
#             for run in paragraph.runs:
#                 run.text = ""
#             # Coloca todo o texto no primeiro run
#             if paragraph.runs:
#                 paragraph.runs[0].text = texto_modificado
#             else:
#                 run = paragraph.add_run(texto_modificado)

#     # Corpo do documento
#     for paragraph in doc.paragraphs:
#         substituir_em_paragrafo(paragraph)

#     # Tabelas
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     substituir_em_paragrafo(paragraph)

#     # Cabeçalhos e rodapés
#     for section in doc.sections:
#         for paragraph in section.header.paragraphs:
#             substituir_em_paragrafo(paragraph)
#         for paragraph in section.footer.paragraphs:
#             substituir_em_paragrafo(paragraph)

#     logger.info("Substituições concluídas")
#     return doc

 
# def gerar_documento(template_path: str, dados_iniciais: Dict, 
#                    impostos: Dict, itens_mt: List[Dict], itens_bt: List[Dict]) -> BytesIO:
#     """
#     Gera o documento Word unificado com os dados fornecidos
#     """
#     logger.info("Iniciando geração de documento unificado")
#     try:
#         # Validar tipos de entrada
#         if not isinstance(template_path, str):
#             raise TypeError("template_path deve ser uma string")
        
#         if not isinstance(dados_iniciais, dict):
#             logger.warning(f"dados_iniciais não é um dicionário. Tipo recebido: {type(dados_iniciais)}")
#             dados_iniciais = {}
        
#         if not isinstance(impostos, dict):
#             logger.warning(f"impostos não é um dicionário. Tipo recebido: {type(impostos)}")
#             impostos = {}
        
#         if not isinstance(itens_mt, list):
#             logger.warning(f"itens_mt não é uma lista. Tipo recebido: {type(itens_mt)}")
#             itens_mt = []
        
#         if not isinstance(itens_bt, list):
#             logger.warning(f"itens_bt não é uma lista. Tipo recebido: {type(itens_bt)}")
#             itens_bt = []

#         # Preparar substituições de texto
#         ips_mt = set(str(item.get('IP', '')) for item in itens_mt if item.get('IP', '00') != '00')
#         ips_bt = set(str(item.get('IP', '')) for item in itens_bt if item.get('IP', '00') != '00')
#         todos_ips = ips_mt.union(ips_bt)

#         # Construir o dicionário de substituições
#         replacements = {
#             '{{CLIENTE}}': str(dados_iniciais.get('cliente', '')),
#             '{{NOMECLIENTE}}': str(dados_iniciais.get('nomeCliente', '')),
#             '{{FONE}}': str(dados_iniciais.get('fone', '')),
#             '{{EMAIL}}': str(dados_iniciais.get('email', '')),
#             '{{BT}}': str(dados_iniciais.get('bt', '')),
#             '{{OBRA}}': str(dados_iniciais.get('obra', ' ')),
#             '{{DIA}}': str(dados_iniciais.get('dia', '')),
#             '{{MES}}': str(dados_iniciais.get('mes', '')),
#             '{{ANO}}': str(dados_iniciais.get('ano', '')),
#             '{{REV}}': str(dados_iniciais.get('rev', '')),
#             '{{LOCAL}}': str(dados_iniciais.get('local_frete', '')),
#             '{{LOCALFRETE}}': str(impostos.get('local_frete_itens', '')),
#             '{{ICMS}}': f"{impostos.get('icms', 0):.1f}%",
#             '{{IP}}': ', '.join(todos_ips) if todos_ips else '',
#             '{obra}': '' if not dados_iniciais.get('obra', '').strip() else 'Obra:'
#         }

#         # Gerar documento
#         buffer = BytesIO()
#         doc = Document(template_path)
        
#         # Substituir texto
#         doc = substituir_texto_documento(doc, replacements)
        
#         # Inserir tabelas sepa
#         doc = inserir_tabelas_separadas(doc, itens_mt, itens_bt, replacements)
        
#         # Salvar documento
#         doc.save(buffer)
#         buffer.seek(0)
        
#         logger.info("Documento unificado gerado com sucesso")
#         return buffer

#     except Exception as e:
#         logger.error(f"Erro ao gerar documento: {str(e)}", exc_info=True)
#         raise
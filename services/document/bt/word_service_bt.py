from .word_tables_bt import create_custom_table, create_custom_table_escopo
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
import os
from docx.shared import Inches
import logging
import streamlit as st

logger=logging.getLogger(__name__)

def verificar_produto_ip(itens_configurados):
   logger = logging.getLogger(__name__)
   logger.info(f"Iniciando verificação de {len(itens_configurados)} itens")
   
   resultados = []
   combinacoes_ja_processadas = set()

   for item in itens_configurados:
       try:
           produto = item.get('Produto', '')
           ip = item.get('IP', '')
           potencia_str = item.get('Potência', '0')
           conector = item.get('Conector', '').upper()
           flange = item.get('Flange', 0)

           # Extrair o valor numérico da potência removendo 'kVA' ou outros caracteres
           try:
               # Remove 'kVA', 'kva', espaços e converte para float
               potencia = float(''.join(c for c in potencia_str if c.isdigit() or c == '.'))
               logger.debug(f"Potência extraída: {potencia} (original: {potencia_str})")
           except (ValueError, AttributeError) as e:
               logger.error(f"Erro ao converter potência '{potencia_str}': {str(e)}")
               potencia = 0

           combinacao = (produto, ip, potencia, conector)
           
           logger.debug(f"""
               Processando item:
               Produto: {produto}
               IP: {ip}
               Potência: {potencia} (original: {potencia_str})
               Conector: {conector}
               Flange: {flange}
           """)

           if combinacao in combinacoes_ja_processadas:
               logger.debug(f"Combinação {combinacao} já processada, pulando...")
               continue

           combinacoes_ja_processadas.add(combinacao)

           # Inicializar variáveis
           caminho_imagem = os.path.join('imagens', produto.lower() if produto else '', 'blank')
           titulo = ''

           # Regras de imagens
           if produto == 'ATT':
               if ip == '00':
                   caminho_imagem = os.path.join('imagens', 'att', 'tg3ip00.png')
               elif ip in ['21', '23']:
                   caminho_imagem = os.path.join('imagens', 'att', f'tg3ip21.png')
               titulo = f"Autotransformador – Modelo TG3 - IP {ip}"
           
           elif produto == 'TM':
               if 2.5 <= potencia <= 20:
                   caminho_imagem = os.path.join('imagens', 'tm', 's3.png')
                   titulo = f"Transformador Monofásico – Modelo S3 - IP {ip}"
               elif 0.05 <= potencia <= 2.5 and 'WAGO' in conector:
                   caminho_imagem = os.path.join('imagens', 'tm', 'b2.png')
                   titulo = f"Transformador Monofásico – Modelo B2 - IP {ip}"
               elif 0.05 <= potencia <= 2.5 and 'SINDAL' in conector:
                   caminho_imagem = os.path.join('imagens', 'tm', 'm9.png')
                   titulo = f"Transformador Monofásico – Modelo M9 – IP {ip}"

           elif produto == 'TT':
               if 0.1 <= potencia <= 2 and 'SINDAL' in conector:
                   caminho_imagem = os.path.join('imagens', 'tt', 'tea2.png')
                   titulo = f"Transformador Trifásico – Modelo TEA2 – IP {ip}"
               elif 0.1 <= potencia <= 2 and 'WAGO' in conector:
                   caminho_imagem = os.path.join('imagens', 'tt', 'ta9.png')
                   titulo = f"Transformador Trifásico – Modelo TA9 – IP {ip}"
               elif 2.5 <= potencia <= 50:
                   if ip == '00':
                       caminho_imagem = os.path.join('imagens', 'tt', 'tg3ip00.png')
                   elif ip in ['21', '23']:
                       caminho_imagem = os.path.join('imagens', 'tt', 'tg3ip21.png')
                   titulo = f"Transformador Trifásico – Modelo TG3 - IP {ip}"
               elif 50 <= potencia <= 360:
                   if ip == '00':
                       caminho_imagem = os.path.join('imagens', 'tt', 'tg12ip00.png')
                   elif ip in ['21', '23']:
                       caminho_imagem = os.path.join('imagens', 'tt', 'tg12ip21.png')
                   elif ip in ['54'] and flange == 1:
                       caminho_imagem = os.path.join('imagens', 'tt', 'tg12ip541flange.png')
                   elif ip in ['54'] and flange == 2:
                       caminho_imagem = os.path.join('imagens', 'tt', 'tg12ip542flange.png')
                   titulo = f"Transformador Trifásico – Modelo TG12 - IP {ip}"

           logger.debug(f"Caminho de imagem definido: {caminho_imagem}")
           logger.debug(f"Título definido: {titulo}")

           resultados.append({
               'Produto': produto,
               'IP': ip,
               'Potência': potencia,
               'Conector': conector,
               'CaminhoImagem': caminho_imagem,
               'Titulo': titulo
           })

       except Exception as e:
           logger.error(f"Erro ao processar item: {str(e)}", exc_info=True)
           continue

   logger.info(f"Processamento concluído. {len(resultados)} resultados gerados")
   return resultados

def inserir_titulo_e_imagem(doc, itens_configurados):
    resultados = verificar_produto_ip(itens_configurados)
    
    # Criar um conjunto para controlar modelos já processados
    modelos_processados = set()
    
    # Função auxiliar para extrair o modelo do título
    def extrair_modelo(titulo):
        if "Modelo" in titulo:
            return titulo.split("Modelo")[1].strip().split("-")[0].strip()
        return ""

    for i, paragraph in enumerate(doc.paragraphs):
        # Encontrar o parágrafo onde as informações devem ser inseridas
        if "A liberação da fabricação ocorrerá imediatamente após o recebimento do pedido de compra." in paragraph.text:
            ultimo_paragrafo = paragraph  # Referência ao ponto inicial de inserção

            for resultado in resultados:
                titulo = resultado['Titulo']
                modelo = extrair_modelo(titulo)
                
                # Verificar se este modelo já foi processado
                if modelo and modelo not in modelos_processados:
                    modelos_processados.add(modelo)
                    
                    caminho_imagem = resultado['CaminhoImagem']
                    
                    logger.debug(f"Inserindo imagem para modelo: {modelo}")
                    logger.debug(f"Título: {titulo}")
                    logger.debug(f"Caminho da imagem: {caminho_imagem}")

                    # Inserir o título
                    p_titulo = doc.add_paragraph(f"{titulo}", style='Heading 2')
                    p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    ultimo_paragrafo._element.addnext(p_titulo._element)
                    ultimo_paragrafo = p_titulo

                    # Inserir a imagem abaixo do título
                    p_imagem = doc.add_paragraph()
                    p_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_imagem = p_imagem.add_run()
                    try:
                        run_imagem.add_picture(caminho_imagem, width=Inches(5.71), height=Inches(5.85))
                        logger.debug(f"Imagem inserida com sucesso para modelo {modelo}")
                    except Exception as e:
                        logger.error(f"Erro ao inserir imagem para modelo {modelo}: {str(e)}")

                    ultimo_paragrafo._element.addnext(p_imagem._element)
                    ultimo_paragrafo = p_imagem
                else:
                    logger.debug(f"Modelo {modelo} já processado, pulando...")

            break  # Parar após adicionar todos os elementos

    logger.info(f"Total de modelos processados: {len(modelos_processados)}")
    logger.debug(f"Modelos processados: {modelos_processados}")

def substituir_texto_documento(doc, replacements):
    def remove_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    # Substituir texto em todos os parágrafos do documento
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                if old_text == "{{IP}}" and not new_text.strip():  # Condicionar a remoção apenas ao campo IP
                    # Remove o parágrafo se o texto de substituição estiver vazio
                    remove_paragraph(paragraph)
                    break  # Sai do loop interno após remover o parágrafo
                else:
                    inline = paragraph.runs
                    for run in inline:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

    # Substituir texto em todas as tabelas do documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            if old_text == "{{IP}}" and not new_text.strip():  # Condicionar a remoção apenas ao campo IP
                                # Remove o parágrafo se o texto de substituição estiver vazio
                                remove_paragraph(paragraph)
                                break  # Sai do loop interno após remover o parágrafo
                            else:
                                inline = paragraph.runs
                                for run in inline:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)

    # Substituir texto no cabeçalho de todas as seções do documento
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    if old_text == "{{IP}}" and not new_text.strip():  # Condicionar a remoção apenas ao campo IP
                        # Remove o parágrafo se o texto de substituição estiver vazio
                        remove_paragraph(paragraph)
                        break  # Sai do loop interno após remover o parágrafo
                    else:
                        inline = paragraph.runs
                        for run in inline:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)


def inserir_tabelas_word(doc, itens_configurados, observacao):
    inserir_titulo_e_imagem(doc, itens_configurados)

    dados_iniciais = st.session_state.get('dados_iniciais', {})
    impostos = st.session_state.get('impostos', {})

    replacements = {
        '{{CLIENTE}}': str(dados_iniciais.get('cliente', '')),
        '{{NOMECLIENTE}}': str(dados_iniciais.get('nomeCliente', '')),
        '{{FONE}}': str(dados_iniciais.get('fone', '')),
        '{{EMAIL}}': str(dados_iniciais.get('email', '')),
        '{{BT}}': str(dados_iniciais.get('bt', '')),
        '{{OBRA}}': str(dados_iniciais.get('obra', ' ')),
        '{{DIA}}': str(dados_iniciais.get('dia', '')),
        '{{MES}}': str(dados_iniciais.get('mes', '')),
        '{{ANO}}': str(dados_iniciais.get('ano', '')),
        '{{REV}}': str(dados_iniciais.get('rev', '')),
        '{{LOCAL}}': str(dados_iniciais.get('local_frete', '')),
        '{{LOCALFRETE}}': str(impostos.get('local_frete', '')),
        '{{ICMS}}': f"{impostos.get('icms', 0):.1f}%",
        '{{IP}}': ', '.join(set(str(item['IP']) for item in itens_configurados 
                              if item['IP'] != '00')),
        '{obra}': '' if not dados_iniciais.get('obra', '').strip() else 'Obra:',
        '{{RESPONSAVEL}}': st.session_state.get('usuario', ''),
        '{{GARANTIA}}': '12',
        '{{VALIDADE}}': '07',
        '{{TRANSPORTE}}': 'O transporte de equipamentos será realizado no formato CIF.' if impostos.get('tipo_frete', '') == 'CIF' else 'O transporte de equipamentos será realizado no formato FOB.',
    }
    import logging
    logging.debug("Iniciando a função 'inserir_tabelas_word'.")

    # Realizar a substituição do texto usando o dicionário replacements
    try:
        substituir_texto_documento(doc, replacements)
        logging.debug("Texto substituído com sucesso usando replacements.")
    except Exception as e:
        logging.error(f"Erro ao substituir texto no documento: {e}")
        raise

    # Procurar o parágrafo "Quadro de Preços" para inserir a tabela logo depois
    try:
        encontrou_quadro_precos = False
        for i, paragraph in enumerate(doc.paragraphs):
            if "Quadro de Preços" in paragraph.text:
                encontrou_quadro_precos = True
                if i + 1 < len(doc.paragraphs):
                      # Certifique-se de que i+1 é válido
                    table = create_custom_table(doc, itens_configurados, observacao)
                    doc.paragraphs[i + 1]._element.addnext(table._element)
                    logging.debug("Tabela de Quadro de Preços inserida com sucesso após o parágrafo.")
                else:
                    logging.warning("Não há parágrafo após 'Quadro de Preços' para inserir a tabela.")
                break
        if not encontrou_quadro_precos:
            logging.warning("Parágrafo 'Quadro de Preços' não encontrado no documento.")
    except Exception as e:
        logging.error(f"Erro ao inserir a tabela de Quadro de Preços: {e}")
        raise

    # Procurar o parágrafo "Escopo de Fornecimento" para inserir a tabela de escopo logo depois
    try:
        encontrou_escopo = False
        for i, paragraph in enumerate(doc.paragraphs):
            if "Escopo de Fornecimento" in paragraph.text:
                encontrou_escopo = True
                if i + 1 < len(doc.paragraphs):  # Certifique-se de que i+1 é válido
                    table_escopo = create_custom_table_escopo(doc, itens_configurados)
                    doc.paragraphs[i + 1]._element.addnext(table_escopo._element)
                    logging.debug("Tabela de Escopo inserida com sucesso após o parágrafo.")
                else:
                    logging.warning("Não há parágrafo após 'Escopo de Fornecimento' para inserir a tabela.")
                break
        if not encontrou_escopo:
            logging.warning("Parágrafo 'Escopo de Fornecimento' não encontrado no documento.")
    except Exception as e:
        logging.error(f"Erro ao inserir a tabela de Escopo: {e}")
        raise

    # Retornar o documento atualizado
    logging.debug("Função 'inserir_tabelas_word' concluída com sucesso.")
    return doc
                           
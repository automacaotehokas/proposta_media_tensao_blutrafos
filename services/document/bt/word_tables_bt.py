from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from .word_formatter_bt import (set_row_height, set_column_widths, 
                              apply_paragraph_formatting, set_cell_shading, 
                              add_double_borders, set_table_left_indent)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import logging


def create_custom_table(doc, itens_configurados, observacao):
    import logging
    
    logging.debug(f"Quantidade de itens configurados: {len(itens_configurados)}")
    num_linhas = len(itens_configurados) + 2  # Uma linha para cada item + cabeçalho + total
    table = doc.add_table(rows=num_linhas, cols=10)  # Adicionando a coluna de IPI

    # Ajustar o alinhamento da tabela para a esquerda
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    table.left_indent = Cm(0)  

    # Set column widths to match MT table
    col_widths = [Cm(1.25), Cm(1.25), Cm(1.5), Cm(2.0), Cm(1.0), Cm(2.5), Cm(1.0), Cm(2.92), Cm(2.92), Cm(1.0)]
    set_column_widths(table, col_widths)

    # Desabilitar o ajuste automático
    table.autofit = False  # Desativa o autofit
    
    # Cabeçalho
    header_row = table.rows[0]
    header_data = ["Item", "Qtde", "Tipo","Potência", "K", "Tensões", "IP",  "Preço Uni. R$", "Preço Total R$", "IPI"]
    for idx, cell in enumerate(header_row.cells):
        if idx < len(header_data):
            cell.text = header_data[idx]
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'Calibri Light (Títulos)'
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # Cor branca
            apply_paragraph_formatting(paragraph, alignment='center')

            # Adicionar sombreamento e bordas duplas
            set_cell_shading(cell, '00543C')  # Cor verde escura
            add_double_borders(cell)
        else:
            logging.warning(f"Índice de célula no cabeçalho fora do intervalo: {idx}")

    set_row_height(header_row, 1)  # 1 cm de altura
    logging.debug("Cabeçalho da tabela configurado com sucesso.")

    # Preenchendo a tabela com os itens configurados
    for idx, item in enumerate(itens_configurados, start=1):
            row = table.rows[idx]  
            row.cells[0].text = str(idx)  
            row.cells[1].text = str(item["Quantidade"])  
            row.cells[2].text = item["Produto"]  # Produto
            row.cells[3].text = item['Potência']
            row.cells[4].text = str(item["Fator K"])  # Fator K
            row.cells[5].text = f"{int(item['Tensão Primária']):d}V /{int(item['Tensão Secundária']):d}V"  
            row.cells[6].text = str(item["IP"])  # IP
            row.cells[7].text = f"R$ {item['Preço Unitário']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")  # Preço unitário
            row.cells[8].text = f"R$ {item['Preço Total']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")  # Preço total
            row.cells[9].text = f"{item.get('IPI', '0')}%"  # IPI

            # Definir a altura da linha como 1,0 cm
            set_row_height(row, 1.0)

            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri Light (Títulos)'
                run.font.size = Pt(11)
                apply_paragraph_formatting(paragraph, alignment='center')

                # Adicionando bordas duplas para cada célula
                add_double_borders(cell)

    # Última linha - Valor Total
    total_row = table.rows[-1]
    if len(total_row.cells) >= 9:  # Verifica se há células suficientes para a linha de total
        total_row.cells[0].merge(total_row.cells[6])  # Mesclar células até "Norma"
        total_row.cells[7].merge(total_row.cells[9])  # Mesclar as colunas Preço Uni., Preço Total e IPI
        total_row.cells[0].text = "Valor Total do Fornecimento:"
        # Convert Preço Total to float before summing
        total = sum(float(item['Preço Total']) for item in itens_configurados)
        total_row.cells[7].text = f"R$ {total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        # Definir altura da linha de total como 0,6 cm
        set_row_height(total_row, 0.6)

        # Adicionar bordas e formatação para a linha do valor total
        for idx in [0, 7]:
            paragraph = total_row.cells[idx].paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'Calibri Light (Títulos)'
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            apply_paragraph_formatting(paragraph, alignment='center', space_before=Pt(0))  # Espaçamento removido
            set_cell_shading(total_row.cells[idx], '00543C')  # Cor de fundo verde escura
            add_double_borders(total_row.cells[idx])
    else:
        logging.warning("Linha de total não possui células suficientes para a mesclagem e preenchimento.")




    logging.debug("Tabela criada com sucesso.")
    return table


def get_ip_text(ip: str, flange: int) -> tuple[str, bool]:
    """
    Retorna o texto do IP formatado com ou sem flanges
    Args:
        ip: valor do IP
        flange: valor do flange
    Returns:
        Tupla com (texto_ip, is_bold)
    """
    ip_text = f"IP-{ip}"
    if flange and flange > 0:
        ip_text += " com flanges"
    return (ip_text, True)

def gerar_escopo_texto(item):
    """
    Gera o texto do escopo com base no tipo de produto, retornando uma lista de tuplas
    onde cada tupla contém (texto, deve_ser_negrito).
    """
    logger = logging.getLogger(__name__)
    logger.debug(f"Iniciando geração de escopo para item: {item}")
    
    try:
        # Extrair valores com tratamento especial para cada campo
        produto = item.get('Produto', '').upper()
        potencia = item.get('Potência', 'N/A')
        tensao_primaria = item.get('Tensão Primária', 'N/A')
        derivacoes_dict = item.get('derivacoes', {})
        taps = derivacoes_dict.get('taps', 'nenhum')
        tensoes_primarias = derivacoes_dict.get('tensoes_primarias', 'nenhum')
        derivacoes = f"{taps}" if taps != "nenhum" else "N/A"
        tensao_secundaria = item.get('Tensão Secundária', 'N/A')
        ip = item.get('IP', 'N/A')
        fator_k = item.get('Fator K', 'N/A')
        material_st = item.get('material', 'N/A')
        material = 'Cobre' if material_st == 'Cu' else 'Alumínio' if material_st == 'Al' else 'N/A'
        
        # Configurações de escopo para diferentes produtos
        if 'ATT' in produto:
            ip_text = f"IP-{item.get('IP', 'N/A')}"
            if item.get('flange', 0) > 0:  # Verifica se tem flange
                ip_text += " com flanges"
            escopo_parts = [
                ("Autotransformador trifásico a seco, classe de tensão 1,1kV, Marca e Fabricação Blutrafos, Potência ", False),
                (potencia, True),
                (", Fator K=", False),
                (str(fator_k), False),
                (", Tensão Primária: ", False),
                (f"{tensao_primaria}V", True),
                (", Derivações: ", False),
                (derivacoes, False),
                (", Tensão Secundária: ", False),
                (f"{tensao_secundaria}V", True),
                (", NBI: N/A, Grupo de Ligação: Yn0, Frequência: 60Hz, Enrolamentos impregnados em verniz a vácuo, "
                 "com resfriamento tipo: AN, Classe de Temperatura materiais isolantes AT/BT: F (155ºC), "
                 f"Elevação Temperatura média dos enrolamentos: 100°C, Materiais dos enrolamentos: {material}, "
                 "Regime de Serviço: Contínuo, Temperatura Ambiente máxima: 40°C, Altitude de Instalação: ≤1000m e grau de proteção ", False),
                (ip_text, True),
                (". Demais características conforme norma ABNT-NBR 5356/11 e acessórios abaixo.", False)
            ]
        elif 'TT' in produto:
            ip_text = f"IP-{item.get('IP', 'N/A')}"
            if item.get('flange', 0) > 0:  # Verifica se tem flange
                ip_text += " com flanges"
            escopo_parts = [
                ("Transformador isolador trifásico a seco, classe de tensão 1,1kV, Marca e Fabricação Blutrafos, Potência ", False),
                (potencia, True),
                (", Fator ", False),
                ("K=1", True),
                (", Tensão Primária: ", False),
                ("380V", True),
                (", Derivações: ", False),
                ("N/A", True),
                (", Tensão Secundária: ", False),
                ("220/127V", True),
                (", NBI: N/A, Grupo de Ligação: Dyn1, Frequência: 60Hz, Enrolamentos impregnados em verniz a vácuo, "
                 "com resfriamento tipo: AN, Classe de Temperatura materiais isolantes AT/BT: F (155ºC), "
                 f"Elevação Temperatura média dos enrolamentos: 100°C, Materiais dos enrolamentos: {material}, "
                 "Regime de Serviço: Contínuo, Temperatura Ambiente máxima: 40°C, Altitude de Instalação: ≤1000m e grau de proteção ", False),
                (ip_text, True),
                (". Demais características conforme norma ", False),
                ("ABNT-NBR 5356/11", True),
                (" e acessórios abaixo.", False)
            ]
        elif 'TM' in produto:
            ip_text = f"IP-{item.get('IP', 'N/A')}"
            if item.get('flange', 0) > 0:  # Verifica se tem flange
                ip_text += " com flanges"
            escopo_parts = [
                ("Transformador isolador monofásico a seco, classe de tensão 1,1kV, Marca e Fabricação Blutrafos, Potência ", False),
                (potencia, True),
                (", Tensão Primária: ", False),
                (f"{tensao_primaria}V", False),
                (", Tensão Secundária: ", False),
                (f"{tensao_secundaria}V", True),
                (", NBI: N/A, Polaridade: Subtrativa, Frequência: 60Hz, Enrolamentos impregnados em verniz a vácuo, "
                 "com resfriamento tipo: AN, Classe de Temperatura materiais isolantes AT/BT: F (155ºC), "
                 f"Elevação Temperatura média dos enrolamentos: 100°C, Materiais dos enrolamentos: {material}, "
                 "Regime de Serviço: Contínuo, Temperatura Ambiente máxima: 40°C, Altitude de Instalação: <1000m, grau de proteção ", False),
                (ip_text, True),
                (", Fator ", False),
                (f"K={fator_k}", True),
                (". Demais características conforme norma ", False),
                ("ABNT-NBR 5356/11", True),
                (" e acessórios abaixo.", False)
            ]
        else:
            logger.error(f"Produto não identificado: {produto}")
            escopo_parts = [("Produto não identificado. Verifique o item e os dados fornecidos.", False)]
        
        return escopo_parts
        
    except Exception as e:
        logger.error(f"Erro ao gerar escopo: {str(e)}", exc_info=True)
        return [("Erro ao gerar escopo do produto. Verifique os dados fornecidos.", False)]

def create_custom_table_escopo(doc, itens_configurados):
    """
    Cria uma tabela de escopo no documento Word com formatação apropriada,
    incluindo texto em negrito onde especificado.
    """
    logger = logging.getLogger(__name__)
    logger.debug(f"Iniciando criação da tabela de escopo com {len(itens_configurados)} itens")
    
    table = doc.add_table(rows=len(itens_configurados) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_left_indent(table, 0)
    table.left_indent = Cm(0)
    table.autofit = False

    # Configurar larguras das colunas
    col_widths = [Cm(1.5), Cm(15.0)]
    set_column_widths(table, col_widths)

    # Configurar cabeçalho
    header_row = table.rows[0]
    header_data = ["Item", "Escopo do Fornecimento:"]
    
    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (T)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center')
        set_row_height(header_row, 1)
        set_cell_shading(cell, '00543C')
        add_double_borders(cell)

    # Preencher a tabela com os itens configurados
    for idx, item in enumerate(itens_configurados, start=1):
        try:
            logger.debug(f"Processando item {idx}: {item}")
            row = table.rows[idx]

            # Configurar coluna "Item"
            row.cells[0].text = str(idx)
            apply_paragraph_formatting(row.cells[0].paragraphs[0], alignment='center')
            add_double_borders(row.cells[0])

            # Gerar e aplicar o texto do escopo com formatação
            escopo_parts = gerar_escopo_texto(item)
            
            # Configurar o parágrafo para o escopo
            escopo_paragraph = row.cells[1].paragraphs[0]
            escopo_paragraph.text = ""  # Limpar o texto existente
            
            # Adicionar cada parte do texto com sua formatação apropriada
            for text, is_bold in escopo_parts:
                run = escopo_paragraph.add_run(text)
                run.font.name = 'Calibri Light (Título)'
                run.font.size = Pt(10)
                run.bold = is_bold

            # Aplicar formatação no parágrafo
            escopo_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph_format = escopo_paragraph.paragraph_format
            paragraph_format.space_after = Pt(2)

            # Adicionar bordas para a célula de escopo
            add_double_borders(row.cells[1])

        except Exception as e:
            logger.error(f"Erro ao processar item {idx}: {str(e)}", exc_info=True)
            row.cells[1].text = "Erro ao gerar escopo. Verifique os dados do item."

    logger.debug("Tabela de escopo criada com sucesso")
    return table
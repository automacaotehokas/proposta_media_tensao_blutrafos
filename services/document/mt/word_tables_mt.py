# services/document/word_tables.py
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .test import formatar_numero_inteiro_ou_decimal
from typing import Dict, List, Any
from .word_formatter_mt import (
    set_row_height, set_column_widths, apply_paragraph_formatting,
    set_cell_shading, add_double_borders, set_table_left_indent
)
import streamlit as st
import logging
logger = logging.getLogger(__name__)

def substituir_texto_documento(doc, replacements):
    def process_run(run, replacements):
        """Processa um único run de texto"""
        modified = False
        text = run.text
        for old_text, new_text in replacements.items():
            if old_text in text:
                text = text.replace(old_text, new_text)
                modified = True
        if modified:
            run.text = text
    def process_paragraph(paragraph):
        """Processa um parágrafo inteiro"""
        if not paragraph or not hasattr(paragraph, 'runs'):
            return
 
        # Guarda o texto original para comparação
        original_text = paragraph.text
        # Primeiro tenta processar run por run
        for run in paragraph.runs:
            process_run(run, replacements)
        # Se ainda existem substituições para fazer, tenta uma abordagem mais agressiva
        if any(old_text in paragraph.text for old_text in replacements):
            try:
                # Preserva o primeiro run com formatação
                first_run = paragraph.runs[0] if paragraph.runs else None
                # Guarda a formatação
                font_props = {}
                if first_run:
                    font = first_run.font
                    font_props = {
                        'name': font.name,
                        'size': font.size,
                        'bold': font.bold,
                        'italic': font.italic,
                        'underline': font.underline,
                        'color': font.color
                    }
                # Limpa todos os runs exceto aqueles com imagens
                text = paragraph.text
                for run in paragraph.runs[:]:
                    if not hasattr(run, '_r'):  # Pula runs inválidos
                        continue
                    try:
                        # Verifica se o run contém uma imagem
                        has_image = bool(run._r.findall("*//pic:pic", {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}))
                        if not has_image:
                            p = run._element
                            p.getparent().remove(p)
                    except Exception:
                        continue
                # Faz todas as substituições no texto
                for old_text, new_text in replacements.items():
                    text = text.replace(old_text, new_text)
                # Adiciona o novo texto como um único run
                new_run = paragraph.add_run(text)
                # Aplica a formatação salva
                if font_props:
                    for prop, value in font_props.items():
                        if value is not None:
                            setattr(new_run.font, prop, value)
            except Exception as e:
                logger.error(f"Erro ao reprocessar parágrafo: {str(e)}")
        # Verifica se precisa remover o parágrafo (caso IP)
        if "{{IP}}" in original_text and replacements.get("{{IP}}", "").strip() == "":
            try:
                p = paragraph._element
                p.getparent().remove(p)
            except Exception as e:
                logger.error(f"Erro ao remover parágrafo: {str(e)}")
 
    # Processa cabeçalhos
        if "{{DIFAL}}" in original_text and (
    replacements.get("{{DIFAL}}", "").strip() in ['0.0', '0', '', '0,0']):
            try:
                p = paragraph._element
                p.getparent().remove(p)
            except Exception as e:
                logger.error(f"Erro ao remover parágrafo: {str(e)}")

 
    # Processa cabeçalhos
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                process_paragraph(paragraph)
    # Processa parágrafos normais
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    # Processa tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    return doc



def determinar_eficiencia(perdas: str) -> str:
    """Determina a eficiência com base nas perdas"""
    if perdas == '5356-D':
        return "D"
    elif perdas == '5356-A':
        return "A"
    elif perdas == '1,2 %':
        return "1,2%"
    elif perdas == '1,0 %':
        return "1%"
    else:
        return "N/A"

def create_custom_table(doc: Document, itens_configurados: List[Dict], observacao: str) -> object:
    """Cria a tabela de preços customizada"""
    num_linhas = len(itens_configurados) + 2  # Uma linha para cada item + cabeçalho + total
    table = doc.add_table(rows=num_linhas, cols=10)  # Adicionando a coluna de IPI

    # Ajustar o alinhamento da tabela para a esquerda
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.left_indent = Cm(0)

    # Definir larguras fixas para as colunas
    col_widths = [Cm(1.1), Cm(1.25), Cm(2.2), Cm(1.0), Cm(2.7), Cm(1.0), 
                  Cm(1.75), Cm(2.63), Cm(2.63), Cm(1.15)]
    
    # Desabilitar o ajuste automático
    table.autofit = False
    set_column_widths(table, col_widths)

    # Cabeçalho
    header_row = table.rows[0]
    header_data = ["Item", "Qtde", "Potência", "K", "Tensões", "IP", "Perda", 
                  "Preço Uni. R$", "Preço Total R$", "IPI"]
    
    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (Títulos)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center')
        set_cell_shading(cell, '00543C')
        add_double_borders(cell)

    set_row_height(header_row, 1)

    # Preenchendo a tabela com os itens
    for idx, item in enumerate(itens_configurados, start=1):
        row = table.rows[idx]
        
        # Formatação da potência
        potencia = item["Potência"]
        potencia_texto = f"{potencia:,.1f}".replace('.', ',') + " kVA" if potencia % 1 != 0 else f"{int(potencia)} kVA"

        # Preenchimento das células
        cells_data = [
            str(idx),
            str(item["Quantidade"]),
            potencia_texto,
            str(item["Fator K"]),
            f"{item['Tensão Primária']}kV/{str(int(item['Tensão Secundária']*0.001)).replace('.',',')}kV",
            str(item["IP"]),
            str(item["Perdas"]),
            f"{item['Preço Unitário']:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            f"{item['Preço Total']:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            f"{item.get('IPI', '0')}%"
        ]

        for cell_idx, cell_text in enumerate(cells_data):
            cell = row.cells[cell_idx]
            cell.text = cell_text
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'Calibri Light (Títulos)'
            run.font.size = Pt(11)
            apply_paragraph_formatting(paragraph, alignment='center')
            add_double_borders(cell)

        set_row_height(row, 1.0)

    # Última linha - Valor Total
    total_row = table.rows[-1]
    total_row.cells[0].merge(total_row.cells[6])
    total_row.cells[7].merge(total_row.cells[9])
    
    # Convert Preço Total to float before summing
    total = sum(float(item['Preço Total']) for item in itens_configurados)
    total_row.cells[0].text = "Valor Total do Fornecimento:"
    total_row.cells[7].text = f"R$ {total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # Formatação da linha de total
    for idx in [0, 7]:
        cell = total_row.cells[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (Títulos)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center', space_before=Pt(0))
        set_cell_shading(cell, '00543C')
        add_double_borders(cell)

    set_row_height(total_row, 0.6)



    return table

def substituir_icms_cirurgico(doc, replacements):
    icms_value = replacements.get("{{ICMS}}", "")
    if not icms_value:
        return

    for paragraph in doc.paragraphs:
        if "{{ICMS}}" not in paragraph.text:
            continue
            
        # Mantém todos os runs originais, só substitui o placeholder
        for run in paragraph.runs:
            if "{{ICMS}}" in run.text:
                # Substitui APENAS o placeholder, mantendo TODO o resto
                original_text = run.text
                before, placeholder, after = original_text.partition("{{ICMS}}")
                
                # Reconstroi o run com o valor novo no meio
                run.text = before + icms_value + after
                
                # Para garantir que não alterou nada além do necessário:
                # 1. Mantém o negrito original do run
                # 2. Não toca em nenhuma outra formatação
                # 3. Não mexe nos runs vizinhos
                break  # Uma única substituição por parágrafo

from docx.shared import Cm,Pt  # Importe a classe Cm para trabalhar com centímetros
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_custom_table_escopo(doc: Document, itens_configurados: List[Dict]) -> object:
    """Cria a tabela de escopo"""
    table = doc.add_table(rows=len(itens_configurados) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_left_indent(table, 0)
    table.left_indent = Cm(0)
    table.autofit = False

    # Definir larguras das colunas
    col_widths = [Cm(1.5), Cm(1.5),Cm(15.0)]
    set_column_widths(table, col_widths)

    # Cabeçalho
    header_row = table.rows[0]
    header_data = ["Item","Qtde", "Escopo do Fornecimento:"]

    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (T)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center')
        set_row_height(header_row, 1)
        set_cell_shading(cell, '00543C')
        add_double_borders(cell)

    # Preenchendo itens
    for idx, item in enumerate(itens_configurados, start=1):
        row = table.rows[idx]
        
        # Coluna "Item"
        row.cells[0].text = str(idx)
        apply_paragraph_formatting(row.cells[0].paragraphs[0], alignment='center')
        add_double_borders(row.cells[0])

        row.cells[1].text = str(item.get("Quantidade", "1"))
        apply_paragraph_formatting(row.cells[1].paragraphs[0], alignment='center')
        add_double_borders(row.cells[1])

        # Preparar dados para o escopo
        classe_tensao = item.get('classe_tensao', '').replace('kV', '').strip()
        qtde= item.get('Quantidade', 'N/A')
        tensao_secundaria = item.get('Tensão Secundária', 'N/A').replace('kV', '').strip()
        eficiencia = determinar_eficiencia(item['Perdas'])
        NBI = item.get('nbi', 'N/A')
        
        # Formatar potência
        potencia = item.get('Potência', 'N/A')
        potencia_formatada = formatar_numero_inteiro_ou_decimal(potencia)


        # Calcular tensão secundária
        tensao_secundaria_str = item.get('Tensão Secundária', '0')
        try:
            tensao_secundaria_float = float(tensao_secundaria_str)
            tensao_calculada = tensao_secundaria_float / 1.73
            tensao_calculada_arredondada = round(tensao_calculada)
            tensao_secundaria_arredondada = round(tensao_secundaria_float)
            tensao_secundaria_texto = f"{tensao_secundaria_arredondada}/{tensao_calculada_arredondada}V"
        except ValueError:
            tensao_secundaria_texto = f"{tensao_secundaria_str}V"

        ip_text = f"IP-{item.get('IP', 'N/A')}"

        # Verificar os flanges selecionados baseado nos acessórios
        flange_at_selecionado = next((acessorio for acessorio in item.get('acessorios', [])
                                    if acessorio['descricao'] == "Flange AT até 15KV (s/ Barramento)"), None)
        flange_bt_selecionado = next((acessorio for acessorio in item.get('acessorios', [])
                                    if acessorio['descricao'] == "Flange BT até 800V (s/ Barramento)"), None)

        if flange_at_selecionado or flange_bt_selecionado:
            ip_text += " com flanges"
        
        escopo_text = (
            f"Transformador Trifásico **isolado a seco**, Classe de tensão **{classe_tensao}/1,1kV**, "
            f"Marca e Fabricação Blutrafos, Potência: **{(potencia_formatada)} kVA**, Fator: **K={item.get('Fator K', 'N/A')}**, "
            f"Tensão **Primária**: **{item.get('Tensão Primária', 'N/A')}kV**, Derivações: **{item.get('Derivações', 'N/A')}**, "
            f"Tensão **Secundária**: **{tensao_secundaria_texto}**, Grupo de Ligação: **Dyn-1**, "
            f"Frequência: **60Hz**, NBI: **{NBI}**, Classe de Temperatura: F (155ºC), "
            f"Elevação Temperatura média dos enrolamentos: **100ºC**, Materiais dos enrolamentos: **Alumínio**, "
            f"Altitude de Instalação: **≤1000m**, Temperatura ambiente máxima: 40°C, "
            f"Alta tensão Encapsulado em Resina Epóxi à Vácuo, Regime de Serviço: Contínuo, "
            f"Tipo de Refrigeração: **AN** e Grau de Proteção: **{ip_text}**, "
            f"Demais características cfe. Norma ABNT-NBR 5356/11 - Eficiência **“{eficiencia}”** e acessórios abaixo."
        )

        # Aplicar o texto com formatação
        escopo_paragraph = row.cells[2].paragraphs[0]
        escopo_paragraph.text = ""
        escopo_parts = escopo_text.split("**")
        
        for i, part in enumerate(escopo_parts):
            run = escopo_paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True
            run.font.name = 'Calibri Light (Título)'
            run.font.size = Pt(10)

        escopo_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        escopo_paragraph.paragraph_format.space_after = Pt(2)
        add_double_borders(row.cells[2])

    return table

def inserir_tabelas_word(doc: Document, itens_configurados: List[Dict], 
                        observacao: str, replacements=None) -> Document:
    """Insere as tabelas no documento e realiza substituições de texto"""
    # Use session state to get initial data
    dados_iniciais = st.session_state.get('dados_iniciais', {})
    impostos = st.session_state.get('impostos', {})

    logger.error("Iniciando rreplaces no documento...")
    # If no replacements provided, create default dictionary
    if replacements is None:
        replacements = {
            '{{CLIENTE}}': str(dados_iniciais.get('cliente', '')),
            '{{NOMECLIENTE}}': str(dados_iniciais.get('nomeCliente', '')),
            '{{FONE}}': str(dados_iniciais.get('fone', '')),
            '{{EMAIL}}': str(dados_iniciais.get('email', '')),
            '{{BT}}': str(dados_iniciais.get('bt', '')),
            '{{OBRA}}': str(dados_iniciais.get('obra', '')),
            '{{DIA}}': str(dados_iniciais.get('dia', '')),
            '{{MES}}': str(dados_iniciais.get('mes', '')),
            '{{ANO}}': str(dados_iniciais.get('ano', '')),
            '{{REV}}': str(dados_iniciais.get('rev', '')),
            '{{LOCAL}}': str(dados_iniciais.get('local_frete', '')),
            '{{LOCALFRETE}}': str(impostos.get('local_frete', '')),
            '{{ICMS}}': f"{formatar_numero_inteiro_ou_decimal(impostos.get('icms', 0))}",
            '{{DIFAL}}': f"{formatar_numero_inteiro_ou_decimal(impostos.get('difal', 0))}" if impostos.get('difal', 0) > 0 else '',
            '{{IP}}': ', '.join(set(str(item['IP']) for item in itens_configurados 
                                  if item['IP'] != '00')),
            '{obra}': 'Obra:' if dados_iniciais.get('obra', '').strip() else '',
            '{{RESPONSAVEL}}': st.session_state.get('usuario', ''),
            '{{GARANTIA}}': '12',
            '{{VALIDADE}}': '07',
                '{{TRANSPORTE}}': f"CIP - {str(dados_iniciais.get('local_frete', ''))}" 
                               if impostos.get('tipo_frete', '') == 'CIP'
                               else f"FOB - {str(dados_iniciais.get('local_frete', ''))}",       
                                 }
        

    # Substituir texto no documento
    substituir_texto_documento(doc, replacements)
    logger.error("Substituições realizadas com sucesso!")

    

    # Inserir tabela de preços
    logger.error("Inserindo tabela de preços...oo")
    for i, paragraph in enumerate(doc.paragraphs):
        if "Quadro de Preços" in paragraph.text:
            table = create_custom_table(doc, itens_configurados, observacao)
            # Ajusta o alinhamento e recuo da tabela
            table.autofit = False
            table.allow_autofit = False
            
            # 3. Ajuste via XML (recuo negativo de 1 cm)
            tbl_pr = table._element.xpath('w:tblPr')
            if tbl_pr:
                tbl_pr = tbl_pr[0]
                tbl_ind = OxmlElement('w:tblInd')
                tbl_ind.set(qn('w:w'), str(int(-1 * 400)))  # -1 cm (567 twips = 1 cm)
                tbl_ind.set(qn('w:type'), 'dxa')
                tbl_pr.append(tbl_ind)
            
            # Insere a tabela após o parágrafo
            doc.paragraphs[i + 1]._element.addnext(table._element)

  
            break

   # Inserir tabela de escopo
    logger.error("Inserindo tabela de escopo...")
    for i, paragraph in enumerate(doc.paragraphs):
        if "Escopo de Fornecimento" in paragraph.text:
            table = create_custom_table_escopo(doc, itens_configurados)
            # Ajusta o alinhamento e recuo da tabela
            table.autofit = False
            table.allow_autofit = False
            
            # 3. Ajuste via XML (recuo negativo de 1 cm)
            tbl_pr = table._element.xpath('w:tblPr')
            if tbl_pr:
                tbl_pr = tbl_pr[0]
                tbl_ind = OxmlElement('w:tblInd')
                tbl_ind.set(qn('w:w'), str(int(-1 * 400)))  # -1 cm (567 twips = 1 cm)
                tbl_ind.set(qn('w:type'), 'dxa')
                tbl_pr.append(tbl_ind)
            
            # Insere a tabela após o parágrafo
            doc.paragraphs[i + 1]._element.addnext(table._element)
            break

    return doc

    



def get_table_width(doc: Document) -> float:
    """Retorna a largura disponível para tabelas no documento"""
    section = doc.sections[0]
    page_width = section.page_width
    margin_left = section.left_margin
    margin_right = section.right_margin
    return page_width - margin_left - margin_right

def format_currency(value: float) -> str:
    """Formata valor monetário no padrão brasileiro"""
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def format_potencia(potencia: Any) -> str:
    """Formata o valor da potência"""
    if isinstance(potencia, (int, float)):
        if potencia % 1 == 0:
            return f"{int(potencia)} kVA"
        return f"{potencia:.1f}".replace('.', ',') + " kVA"
    return str(potencia)

def calcular_tensao_secundaria(tensao_str: str) -> str:
    """Calcula e formata a tensão secundária"""
    try:
        tensao = float(tensao_str)
        tensao_fase = tensao / 1.73
        return f"{round(tensao)}/{round(tensao_fase)}V"
    except (ValueError, TypeError):
        return f"{tensao_str}V"

# def formatar_numero_inteiro_ou_decimal(valor):
#     """
#     Converte um número para inteiro se for um número inteiro,
#     ou mantém uma casa decimal substituindo . por , se tiver casa decimal.
    
#     Exemplos:
#     1000.00 -> 1000
#     125.5 -> 125,5
#     125.50 -> 125
#     """
#     try:
#         # Converte para float para garantir formato correto
#         valor_float = float(valor)
        
#         # Se for um número inteiro
#         if valor_float.is_integer():
#             return str(int(valor_float))
        
#         # Se tiver casa decimal
#         return f"{valor_float:.1f}".replace('.', ',')
    
#     except (ValueError, TypeError):
#         return str(valor)

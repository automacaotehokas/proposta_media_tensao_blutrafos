from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH 
import math
from docx.enum.table import WD_TABLE_ALIGNMENT



# Função para definir a altura de uma linha
def set_row_height(row, height_cm):
    tr = row._tr
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # Multiplicando por 567 para converter cm para twips
    trHeight.set(qn('w:hRule'), 'exact')  # Define altura exata
    tr.append(trHeight)

# Função para aplicar largura exata a uma célula
def set_column_widths(table, col_widths):
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

# Função para aplicar formatação a um parágrafo
def apply_paragraph_formatting(paragraph, alignment='center', space_before=Pt(5), space_after=Pt(0)):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = space_before
    paragraph_format.space_after = space_after
    paragraph_format.line_spacing = 1  # Espaçamento simples
    paragraph.alignment = {'left': 0, 'center': 1, 'right': 2}.get(alignment, 1)

# Função para aplicar sombreamento a uma célula
def set_cell_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

# Função para adicionar bordas duplas a uma célula
def add_double_borders(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'double')  # Bordas duplas
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

# Função para determinar a eficiência com base nas perdas
def determinar_eficiencia(perdas):
    if perdas == '5356-D':
        return "D"
    elif perdas == '5356-A':
        return "A"
    elif perdas == '1,2 %':
        return "1,2%"
    elif perdas == '1,0 %':
        return "1%"
    else:
        return "N/A"  # Valor padrão se não for encontrado

# Função para criar a tabela do Quadro de Preços
def create_custom_table(doc, itens_configurados, observacao):
    num_linhas = len(itens_configurados) + 2  # Uma linha para cada item + cabeçalho + total
    table = doc.add_table(rows=num_linhas, cols=10)  # Adicionando a coluna de IPI

    # Ajustar o alinhamento da tabela para a esquerda
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Definir a indentação esquerda da tabela como zero
    table.left_indent = Cm(0)  # Adicionando a coluna de IPI

    # Definir larguras fixas para as colunas
    col_widths = [Cm(1.1), Cm(1.25), Cm(2.2), Cm(1.0), Cm(2.7), Cm(1.0), Cm(1.75), Cm(2.63), Cm(2.63), Cm(1.15)]

    # Desabilitar o ajuste automático
    table.autofit = False  # Desativa o autofit

    # Aplicar largura das colunas
    set_column_widths(table, col_widths)

    # Cabeçalho
    header_row = table.rows[0]
    header_data = ["Item", "Qtde", "Potência", "K", "Tensões", "IP", "Perda", "Preço Uni. R$", "Preço Total R$", "IPI"]
    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (Títulos)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # Cor branca
        apply_paragraph_formatting(paragraph, alignment='center')

        # Adicionar sombreamento e bordas duplas
        set_cell_shading(cell, '00543C')  # Cor verde escura
        add_double_borders(cell)

    set_row_height(header_row, 1)  # 1 cm de altura

    # Preenchendo a tabela com os itens configurados
    for idx, item in enumerate(itens_configurados, start=1):
        row = table.rows[idx]
        row.cells[0].text = str(idx)  # Número do item
        row.cells[1].text = str(item["Quantidade"])  # Quantidade
        row.cells[2].text = f"{item['Potência']:.1f}".replace('.', ',') + " kVA" if item["Potência"] % 1 != 0 else f"{int(item['Potência'])} kVA"
        row.cells[3].text = str(item["Fator K"])  # Fator K
        row.cells[4].text = f"{item['Tensão Primária']}kV /{item['Tensão Secundária']} V"  # Tensão
        row.cells[5].text = str(item["IP"])  # IP
        row.cells[6].text = str(item["Perdas"])  # Norma (antiga Perda)
        row.cells[7].text = f"{item['Preço Unitário']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")  # Preço unitário
        row.cells[8].text = f"{item['Preço Total']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")  # Preço total
        row.cells[9].text = f"{item.get('IPI', '0')}%"  # IPI

        # Definir a altura da linha como 1,0 cm
        set_row_height(row, 1.0)

        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'Calibri Light (Títulos)'
            run.font.size = Pt(11)
            apply_paragraph_formatting(paragraph, alignment='center')

            # Adicionando bordas duplas para cada célula
            add_double_borders(cell)

    # Última linha - Valor Total
    total_row = table.rows[-1]
    total_row.cells[0].merge(total_row.cells[6])  # Mesclar células até "Norma"
    total_row.cells[7].merge(total_row.cells[9])  # Mesclar as colunas Preço Uni., Preço Total e IPI
    total_row.cells[0].text = "Valor Total do Fornecimento:"
    total_row.cells[7].text = f"R$ {sum(item['Preço Total'] for item in itens_configurados):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # Definir altura da linha de total como 0,6 cm
    set_row_height(total_row, 0.6)

    # Adicionar bordas e formatação para a linha do valor total
    for idx in [0, 7]:
        paragraph = total_row.cells[idx].paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (Títulos)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center', space_before=Pt(0))  # Espaçamento removido
        set_cell_shading(total_row.cells[idx], '00543C')  # Cor de fundo verde escura
        add_double_borders(total_row.cells[idx])

    # Adicionando a linha de observação
    obs_row = table.add_row()
    obs_row.cells[0].merge(obs_row.cells[9])  # Mesclar todas as células da linha de observação
    obs_cell = obs_row.cells[0]
    obs_cell.text = f"Obs.: {observacao}"
    obs_paragraph = obs_cell.paragraphs[0]
    obs_run = obs_paragraph.runs[0]
    obs_run.font.name = 'Calibri Light (Títulos)'
    obs_run.font.size = Pt(11)
    apply_paragraph_formatting(obs_paragraph, alignment='left', space_before=Pt(0))  # Espaçamento removido

    # Adicionar bordas duplas à célula de observação
    add_double_borders(obs_cell)

    return table


def set_table_left_indent(table, indent):
    tbl_pr = table._tbl.tblPr
    tbl_indent = tbl_pr.xpath("w:tblInd")
    if tbl_indent:
        tbl_indent[0].set(qn('w:w'), str(indent))
        tbl_indent[0].set(qn('w:type'), 'dxa')
    else:
        tbl_indent_element = OxmlElement('w:tblInd')
        tbl_indent_element.set(qn('w:w'), str(indent))
        tbl_indent_element.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_indent_element)


def create_custom_table_escopo(doc, itens_configurados):
    table = doc.add_table(rows=len(itens_configurados) + 1, cols=2)  # Tabela com 2 colunas (Item e Escopo do Fornecimento)

    # Ajustar o alinhamento da tabela para a esquerda
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    set_table_left_indent(table, 0)

    # Definir a indentação esquerda da tabela como zero
    table.left_indent = Cm(0) # Tabela com 2 colunas (Item e Escopo do Fornecimento)

    # Desabilitar o ajuste automático
    table.autofit = False  # Desativa o autofit

    # Definir larguras fixas para as colunas
    col_widths = [Cm(1.5), Cm(15.0)]  # Ajuste as larguras das colunas para não exceder o tamanho total
    set_column_widths(table, col_widths)

    # Cabeçalho
    header_row = table.rows[0]
    header_data = ["Item", "Escopo do Fornecimento:"]

    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light (T)'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # Cor branca
        apply_paragraph_formatting(paragraph, alignment='center')

        # Definindo altura do cabeçalho como 1 cm
        set_row_height(header_row, 1)

        # Adicionando bordas duplas e cor de fundo
        set_cell_shading(cell, '00543C')  # Cor verde escura
        add_double_borders(cell)

    # Preenchendo a tabela com os itens configurados
    for idx, item in enumerate(itens_configurados, start=1):
        row = table.rows[idx]

        # Coluna "Item"
        row.cells[0].text = str(idx)  # Número do item
        apply_paragraph_formatting(row.cells[0].paragraphs[0], alignment='center')
        add_double_borders(row.cells[0])

        # Verificar se a chave 'Classe de Tensao' existe, senão usar um valor padrão ou ajustar a chave correta
        classe_tensao = item.get('classe_tensao', 'N/A').replace('kV', '').strip()  # Remove 'kV'
        tensao_secundaria = item.get('Tensão Secundária', 'N/A').replace('kV', '').strip()  # Tensão Secundária
        eficiencia = determinar_eficiencia(item['Perdas'])

        potencia = item.get('Potência', 'N/A')
        if isinstance(potencia, (int, float)):
            if potencia % 1 == 0:  # Verifica se a parte decimal é 0
                potencia_formatada = f"{int(potencia)} kVA"  # Converte para inteiro
            else:
                potencia_formatada = f"{potencia:.1f}".replace('.', ',') + " kVA"  # Formata com uma casa decimal e troca ponto por vírgula
        else:
            potencia_formatada = potencia 

        # Pegando o valor da Tensão Secundária como texto e tentando convertê-lo para float
        tensao_secundaria_str = item.get('Tensão Secundária', '0')

        try:
            # Tenta converter a tensão secundária para float
            # Convertendo a tensão secundária para float
            tensao_secundaria_float = float(tensao_secundaria_str)

            # Cálculo correto da tensão secundária a partir da tensão de linha (usando sqrt(3))
            raiz_tensao_secundaria = 1.73
            tensao_calculada = tensao_secundaria_float / raiz_tensao_secundaria

            # Forçando o arredondamento para cima (próximo número inteiro)
            tensao_calculada_arredondada = round(tensao_calculada)  # Arredonda sempre para cima
            tensao_secundaria_arredondada = round(tensao_secundaria_float)  # Arredonda sempre para cima

            # Formata o texto com os dois valores arredondados
            tensao_secundaria_texto = f"{tensao_secundaria_arredondada}/{tensao_calculada_arredondada}V"
        
        except ValueError:
            # Se a conversão falhar, exibe apenas o valor original como texto e um aviso
            tensao_secundaria_texto = f"{tensao_secundaria_str}V (valor inválido para cálculo)"

        # Texto de escopo com negrito em palavras selecionadas
        escopo_text = (
            f"Transformador Trifásico **isolado a seco**, Classe de tensão **{classe_tensao}/1,1kV**, "
            f"Marca e Fabricação Blutrafos, Potência: **{potencia_formatada}**, Fator: **K={item.get('Fator K', 'N/A')}**, "
            f"Tensão **Primária**: **{item.get('Tensão Primária', 'N/A')}kV**, Derivações: **{item.get('Derivações', 'N/A')}**, "
            f"Tensão **Secundária**: **{tensao_secundaria_texto}**, Grupo de Ligação: **Dyn-1**, "
            f"Frequência: **60Hz**, NBI: **95kV**, Classe de Temperatura: F (155ºC), "
            f"Elevação Temperatura média dos enrolamentos: **100ºC**, Materiais dos enrolamentos: **Alumínio**, "
            f"Altitude de Instalação: **≤1000m**, Temperatura ambiente máxima: 40°C, "
            f"Alta tensão Encapsulado em Resina Epóxi à Vácuo, Regime de Serviço: Contínuo, "
            f"Tipo de Refrigeração: **AN** e Grau de Proteção: **IP-{item.get('IP', 'N/A')}**, "
            f"Demais características cfe. Norma ABNT-NBR 5356/11 - Eficiência **“{eficiencia}”** e acessórios abaixo."
        )

        # Configurar o parágrafo e aplicar negrito no texto configurado
        escopo_paragraph = row.cells[1].paragraphs[0]
        escopo_paragraph.text = ""  # Limpar o texto padrão
        escopo_parts = escopo_text.split("**")  # Separar o texto nas partes que precisam de negrito

        for i, part in enumerate(escopo_parts):
            run = escopo_paragraph.add_run(part)
            if i % 2 == 1:  # Partes ímpares serão negritadas
                run.bold = True
            run.font.name = 'Calibri Light (Título)'
            run.font.size = Pt(10)

        # Aplicar formatação no parágrafo
        escopo_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alinhamento justificado
        paragraph_format = escopo_paragraph.paragraph_format
        paragraph_format.space_after = Pt(2)

        # Bordas para a célula de escopo
        add_double_borders(row.cells[1])

    return table


# Função para substituir texto no documento, inclusive nas tabelas e no cabeçalho
# Função para substituir texto no documento, inclusive nas tabelas e no cabeçalho
def substituir_texto_documento(doc, replacements):
    def remove_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    # Substituir texto em todos os parágrafos do documento
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                if old_text == "{{IP}}" and not new_text.strip():  # Condicionar a remoção apenas ao campo IP
                    # Remove o parágrafo se o texto de substituição estiver vazio
                    remove_paragraph(paragraph)
                    break  # Sai do loop interno após remover o parágrafo
                else:
                    inline = paragraph.runs
                    for run in inline:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

    # Substituir texto em todas as tabelas do documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            if old_text == "{{IP}}" and not new_text.strip():  # Condicionar a remoção apenas ao campo IP
                                # Remove o parágrafo se o texto de substituição estiver vazio
                                remove_paragraph(paragraph)
                                break  # Sai do loop interno após remover o parágrafo
                            else:
                                inline = paragraph.runs
                                for run in inline:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)

    # Substituir texto no cabeçalho de todas as seções do documento
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    if old_text == "{{IP}}" and not new_text.strip():  # Condicionar a remoção apenas ao campo IP
                        # Remove o parágrafo se o texto de substituição estiver vazio
                        remove_paragraph(paragraph)
                        break  # Sai do loop interno após remover o parágrafo
                    else:
                        inline = paragraph.runs
                        for run in inline:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)


# Função para inserir as tabelas e realizar substituições de texto
def inserir_tabelas_word(doc, itens_configurados, observacao, replacements):
    # Realizar a substituição do texto usando o dicionário replacements
    substituir_texto_documento(doc, replacements)

    # Procurar o parágrafo "Quadro de Preços" para inserir a tabela logo depois
    for i, paragraph in enumerate(doc.paragraphs):
        if "Quadro de Preços" in paragraph.text:
            # Inserir a tabela de Quadro de Preços logo após o parágrafo
            table = create_custom_table(doc, itens_configurados, observacao)
            doc.paragraphs[i+1]._element.addnext(table._element)
            break

    # Procurar o parágrafo "Escopo de Fornecimento" para inserir a tabela de escopo logo depois
    for i, paragraph in enumerate(doc.paragraphs):
        if "Escopo de Fornecimento" in paragraph.text:
            # Inserir a tabela de escopo logo após o parágrafo
            table_escopo = create_custom_table_escopo(doc, itens_configurados)
            doc.paragraphs[i+1]._element.addnext(table_escopo._element)
            break

    # Retornar o documento atualizado
    return doc
    # Salvar o documento no caminho de saída
    doc.save(output_path)

